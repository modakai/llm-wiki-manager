from dataclasses import dataclass
from pathlib import Path


class ParserError(RuntimeError):
    """文档解析失败，向 API 层提供可理解的错误摘要。"""


@dataclass(frozen=True)
class ParsedDocument:
    """解析后的纯文本结果。"""

    text: str
    source_type: str


def parse_document(path: Path) -> ParsedDocument:
    """根据扩展名选择解析器，MVP 不处理 OCR 和老式 doc 转换。"""

    suffix = path.suffix.lower().lstrip(".")
    if suffix in {"txt", "md"}:
        text = _read_text_file(path)
    elif suffix == "pdf":
        text = _read_pdf(path)
    elif suffix == "docx":
        text = _read_docx(path)
    elif suffix == "doc":
        raise ParserError("暂不直接解析 .doc，请安装 LibreOffice 后转换为 .docx 再导入。")
    else:
        raise ParserError(f"不支持的文件类型：.{suffix}")

    cleaned = text.strip()
    if not cleaned:
        raise ParserError("没有提取到有效文本，扫描版 PDF 需要 OCR，暂不在 MVP 支持。")
    return ParsedDocument(text=cleaned, source_type=suffix)


def _read_text_file(path: Path) -> str:
    """读取常见文本编码，优先使用 UTF-8。"""

    for encoding in ("utf-8", "utf-8-sig", "gb18030"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise ParserError("文本编码无法识别，请转换为 UTF-8 后重试。")


def _read_pdf(path: Path) -> str:
    """提取文本型 PDF 内容；扫描版 PDF 会得到空文本。"""

    try:
        import fitz
    except ImportError as exc:
        raise ParserError("缺少 PyMuPDF，无法解析 PDF。") from exc

    try:
        with fitz.open(path) as document:
            return "\n".join(page.get_text("text") for page in document)
    except Exception as exc:  # noqa: BLE001 - 第三方解析异常需要统一转换。
        raise ParserError(f"PDF 解析失败：{exc}") from exc


def _read_docx(path: Path) -> str:
    """提取 docx 段落文本，不解析图片和复杂版式。"""

    try:
        from docx import Document
    except ImportError as exc:
        raise ParserError("缺少 python-docx，无法解析 DOCX。") from exc

    try:
        document = Document(path)
        return "\n".join(paragraph.text for paragraph in document.paragraphs)
    except Exception as exc:  # noqa: BLE001 - 第三方解析异常需要统一转换。
        raise ParserError(f"DOCX 解析失败：{exc}") from exc
